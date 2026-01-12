#!/usr/bin/env python3
"""
Batch Quotation Generator
Reads CSV file and generates .docx and .pdf quotations
Usage: python3 quotation_batch.py <path_to_csv_file>
"""

import sys
import csv
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess

# Constants
COMPANY_NAME = "Cebu Best Value Trading Corp."
COMPANY_LOCATION = "Cebu City"
COMPANY_PHONE = "032-2670573"
COMPANY_MOBILE_SUN = "09325314857"
COMPANY_MOBILE_GLOBE = "09154657503"
SERVICES = "Sales    Installation    Service    Repair\nDuctworks"
OUTPUT_DIR = Path(__file__).parent.parent / "data" / "output"

def parse_csv(csv_path):
    """Parse CSV file into header and items."""
    with open(csv_path, 'r') as f:
        lines = f.readlines()

    # Find header and items sections
    header_section = None
    items_section = None

    header_start = None
    items_start = None

    for i, line in enumerate(lines):
        if line.strip() == "[ITEMS]":
            items_start = i
            break
        if header_start is None and not line.startswith('['):
            header_start = i

    if header_start is not None and items_start is not None:
        header_section = lines[header_start:items_start]
        items_section = lines[items_start + 1:]
    elif header_start is not None:
        header_section = lines[header_start:]
    else:
        raise ValueError("CSV format invalid. Expected [HEADER] or [ITEMS] markers.")

    # Parse header
    header = {}
    if header_section:
        reader = csv.DictReader(header_section)
        header_row = next(reader)
        header = {k: (v.strip() if v else None) for k, v in header_row.items()}

    # Parse items
    items = []
    if items_section:
        reader = csv.DictReader(items_section)
        current_item = None
        for row in reader:
            if not row.get('item_name') or not row['item_name'].strip():
                continue

            item_name = row['item_name'].strip()
            task_name = row.get('task_name', '').strip()
            task_cost = float(row.get('task_cost', 0)) if row.get('task_cost', '').strip() else 0
            quantity = float(row.get('quantity', 1)) if row.get('quantity', '').strip() else 1
            item_warranty = row.get('item_warranty', '').strip() or None

            # Check if new item or continuation
            if current_item is None or current_item['name'] != item_name:
                # Save previous item
                if current_item:
                    items.append(current_item)

                # Start new item
                current_item = {
                    'name': item_name,
                    'ac_brand': row.get('ac_brand', '').strip() or None,
                    'ac_model': row.get('ac_model', '').strip() or None,
                    'warranty': item_warranty,
                    'tasks': [],
                    'item_total': 0,
                }

            # Add task
            current_item['tasks'].append({
                'name': task_name,
                'cost': task_cost,
                'quantity': quantity,
            })
            current_item['item_total'] += task_cost * quantity

        # Add last item
        if current_item:
            items.append(current_item)

    return header, items

def add_horizontal_line(doc):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def create_docx(header, items):
    """Create .docx file from parsed data."""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Professional Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run(COMPANY_NAME)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Calibri'

    # Location
    location_para = doc.add_paragraph()
    location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location_para.paragraph_format.space_after = Pt(3)
    location_run = location_para.add_run(f"{COMPANY_LOCATION}")
    location_run.font.size = Pt(11)
    location_run.font.name = 'Calibri'

    # Horizontal line separator
    add_horizontal_line(doc)

    # Contact info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(2)
    contact = contact_para.add_run(f"Telephone: {COMPANY_PHONE} | Mobile: Sun-{COMPANY_MOBILE_SUN} | Globe-{COMPANY_MOBILE_GLOBE}")
    contact.font.size = Pt(9)
    contact.font.name = 'Calibri'

    # Services
    services_para = doc.add_paragraph()
    services_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    services_para.paragraph_format.space_after = Pt(12)
    services_run = services_para.add_run(SERVICES)
    services_run.font.size = Pt(10)
    services_run.font.name = 'Calibri'

    # Customer info section
    date_str = header.get('date', datetime.today().strftime("%m/%d/%Y"))

    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(3)
    date_label = date_para.add_run("Date: ")
    date_label.bold = True
    date_label.font.name = 'Calibri'
    date_val = date_para.add_run(date_str)
    date_val.font.name = 'Calibri'

    to_para = doc.add_paragraph()
    to_para.paragraph_format.space_after = Pt(2)
    to_label = to_para.add_run("To: ")
    to_label.bold = True
    to_label.font.name = 'Calibri'
    to_val = to_para.add_run(header.get('customer_name', ''))
    to_val.font.name = 'Calibri'

    if header.get('customer_location'):
        loc_para = doc.add_paragraph(header['customer_location'])
        loc_para.paragraph_format.space_after = Pt(2)
        loc_para.paragraph_format.left_indent = Inches(0.25)
        for run in loc_para.runs:
            run.font.name = 'Calibri'

    attention_text = header.get('attention', '')
    if header.get('phone'):
        attention_text += f" | {header['phone']}"
    if attention_text:
        att_para = doc.add_paragraph()
        att_para.paragraph_format.space_after = Pt(8)
        att_label = att_para.add_run("Attention: ")
        att_label.bold = True
        att_label.font.name = 'Calibri'
        att_val = att_para.add_run(attention_text)
        att_val.font.name = 'Calibri'

    doc.add_paragraph()  # Spacing

    # Salutation
    salutation = doc.add_paragraph("Sir/Madame,")
    salutation.paragraph_format.space_after = Pt(8)
    for run in salutation.runs:
        run.font.name = 'Calibri'

    # Opening
    opening = doc.add_paragraph()
    opening.paragraph_format.space_after = Pt(8)
    opening_text = ("This is with reference to your request for quotation for the installation/repair of the "
                    "following air-conditioner to be installed at:")
    opening_run = opening.add_run(opening_text)
    opening_run.font.name = 'Calibri'
    if header.get('installation_location'):
        loc_run = opening.add_run(f" {header['installation_location']}")
        loc_run.font.name = 'Calibri'

    pleased = doc.add_paragraph("We are pleased to quote the following:")
    pleased.paragraph_format.space_after = Pt(8)
    for run in pleased.runs:
        run.font.name = 'Calibri'

    # Items section
    doc_type = header.get('doc_type', 'summary')
    heading_text = "Summary of Quotations" if doc_type == "summary" else "Job to be done:"
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    heading_run = heading.add_run(heading_text)
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    heading_run.font.name = 'Calibri'

    grand_total = 0
    for i, item in enumerate(items, 1):
        # Item header
        item_header = doc.add_paragraph()
        item_header.paragraph_format.space_before = Pt(6)
        item_header.paragraph_format.space_after = Pt(3)
        item_header_run = item_header.add_run(f"{i}. {item['name']}")
        item_header_run.bold = True
        item_header_run.font.size = Pt(11)
        item_header_run.font.name = 'Calibri'
        if item.get('ac_brand'):
            brand_run = item_header.add_run(f" - {item['ac_brand']}")
            brand_run.font.name = 'Calibri'
        if item.get('ac_model'):
            model_run = item_header.add_run(f" {item['ac_model']}")
            model_run.font.name = 'Calibri'

        # Tasks
        for j, task in enumerate(item['tasks'], 1):
            task_para = doc.add_paragraph()
            task_para.paragraph_format.left_indent = Inches(0.25)
            task_para.paragraph_format.space_after = Pt(2)
            task_num = task_para.add_run(f"{j}. ")
            task_num.font.name = 'Calibri'
            task_name = task_para.add_run(f"{task['name']}")
            task_name.font.name = 'Calibri'
            if task['cost'] > 0:
                cost_text = task_para.add_run(f" – ₱{task['cost']:,.0f}")
                cost_text.font.name = 'Calibri'
            if task['quantity'] > 1:
                qty_text = task_para.add_run(f" × {task['quantity']} = ₱{task['cost'] * task['quantity']:,.0f}")
                qty_text.font.name = 'Calibri'

        # Item total
        total_para = doc.add_paragraph()
        total_para.paragraph_format.left_indent = Inches(0.25)
        total_para.paragraph_format.space_after = Pt(3)
        total_run = total_para.add_run(f"Total Price – ₱ {item['item_total']:,.2f}")
        total_run.bold = True
        total_run.font.name = 'Calibri'

        # Item warranty
        if item.get('warranty'):
            warranty_para = doc.add_paragraph()
            warranty_para.paragraph_format.left_indent = Inches(0.25)
            warranty_para.paragraph_format.space_after = Pt(8)
            warranty_run = warranty_para.add_run(f"Warranty: {item['warranty']}")
            warranty_run.font.name = 'Calibri'

        grand_total += item['item_total']

    # Grand total
    if len(items) > 1:
        grand_para = doc.add_paragraph()
        grand_para.paragraph_format.space_before = Pt(8)
        grand_para.paragraph_format.space_after = Pt(12)
        grand_run = grand_para.add_run(f"Total Price of All Items – ₱ {grand_total:,.2f}")
        grand_run.bold = True
        grand_run.font.name = 'Calibri'

    # Summary info section
    add_horizontal_line(doc)

    if header.get('note'):
        note_para = doc.add_paragraph()
        note_para.paragraph_format.space_after = Pt(6)
        note_label = note_para.add_run("Note: ")
        note_label.bold = True
        note_label.font.name = 'Calibri'
        note_val = note_para.add_run(header['note'])
        note_val.font.name = 'Calibri'

    if header.get('payment'):
        payment_para = doc.add_paragraph()
        payment_para.paragraph_format.space_after = Pt(3)
        payment_label = payment_para.add_run("Terms of Payment: ")
        payment_label.bold = True
        payment_label.font.name = 'Calibri'
        payment_val = payment_para.add_run(header['payment'])
        payment_val.font.name = 'Calibri'

    if header.get('warranty'):
        warranty_para = doc.add_paragraph()
        warranty_para.paragraph_format.space_after = Pt(3)
        warranty_label = warranty_para.add_run("Warranty: ")
        warranty_label.bold = True
        warranty_label.font.name = 'Calibri'
        warranty_val = warranty_para.add_run(header['warranty'])
        warranty_val.font.name = 'Calibri'

    if header.get('exceptions'):
        excep_para = doc.add_paragraph()
        excep_para.paragraph_format.space_after = Pt(12)
        excep_label = excep_para.add_run("Exception: ")
        excep_label.bold = True
        excep_label.font.name = 'Calibri'
        excep_val = excep_para.add_run(header['exceptions'])
        excep_val.font.name = 'Calibri'

    # Closing
    thanks_para = doc.add_paragraph()
    thanks_para.paragraph_format.space_before = Pt(6)
    thanks_para.paragraph_format.space_after = Pt(12)
    thanks_run = thanks_para.add_run("Thank you very much for giving us the opportunity to quote and we hope to have the pleasure of serving you.")
    thanks_run.font.name = 'Calibri'

    doc.add_paragraph()  # Spacing

    closing = doc.add_paragraph("Very Truly Yours,")
    closing.paragraph_format.space_after = Pt(40)
    for run in closing.runs:
        run.font.name = 'Calibri'

    # Signature block
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = False
    manager = header.get('manager', 'J.B Yap Jr.')

    # Conforme line
    conforme_cell = sig_table.rows[0].cells[0]
    conforme_para = conforme_cell.paragraphs[0]
    conforme_run = conforme_para.add_run("Conforme:_______________")
    conforme_run.font.name = 'Calibri'

    # Manager name cell
    manager_cell = sig_table.rows[0].cells[1]
    manager_para = manager_cell.paragraphs[0]
    manager_name_run = manager_para.add_run(manager)
    manager_name_run.font.name = 'Calibri'
    manager_name_run.bold = True
    manager_para.add_run("\n")
    manager_title_run = manager_para.add_run("Manager")
    manager_title_run.font.name = 'Calibri'

    # Date line
    date_cell = sig_table.rows[1].cells[0]
    date_para = date_cell.paragraphs[0]
    date_run = date_para.add_run("Date:_______________")
    date_run.font.name = 'Calibri'

    # Empty cell for second column
    sig_table.rows[1].cells[1].text = ""

    return doc

def save_and_convert(doc, filename):
    """Save .docx and convert to .pdf."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    docx_path = OUTPUT_DIR / f"{filename}.docx"
    pdf_path = OUTPUT_DIR / f"{filename}.pdf"

    # Save docx
    doc.save(docx_path)
    print(f"  ✓ Saved: {docx_path}")

    # Convert to PDF
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTPUT_DIR), str(docx_path)],
            check=True,
            capture_output=True,
            timeout=60
        )
        print(f"  ✓ Converted: {pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"  ⚠ PDF conversion failed: {e.stderr.decode() if e.stderr else 'Unknown error'}")
    except FileNotFoundError:
        print("  ⚠ LibreOffice not found. PDF conversion skipped.")
    except subprocess.TimeoutExpired:
        print("  ⚠ PDF conversion timeout.")

def main():
    """Main execution flow."""
    if len(sys.argv) < 2:
        print("Usage: python3 quotation_batch.py <path_to_csv_file>")
        sys.exit(1)

    csv_path = Path(sys.argv[1])

    if not csv_path.exists():
        print(f"Error: File not found: {csv_path}")
        sys.exit(1)

    print("\n" + "=" * 60)
    print("  BATCH QUOTATION GENERATOR")
    print("=" * 60)

    try:
        print(f"\n  Reading: {csv_path}")
        header, items = parse_csv(csv_path)

        customer = header.get('customer_name', 'Unknown')
        date = header.get('date', datetime.today().strftime("%m/%d/%Y"))
        item_count = len(items)
        total = sum(item['item_total'] for item in items)

        print(f"  ✓ Parsed successfully")
        print(f"    Customer: {customer}")
        print(f"    Date: {date}")
        print(f"    Items: {item_count}")
        print(f"    Total: ₱{total:,.2f}\n")

        print("  Generating documents...")
        doc = create_docx(header, items)
        filename = f"{customer}-{date.replace('/', '_')}"
        save_and_convert(doc, filename)

        print(f"\n  ✓ Complete. Files saved to: {OUTPUT_DIR}\n")

    except Exception as e:
        print(f"\n✗ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
