#!/usr/bin/env python3
"""
Sequential CLI Quotation Generator
Generates .docx and .pdf quotations for Cebu Best Value Trading Corp.
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess

# Constants
COMPANY_NAME = "Cebu Best Value Trading Corp."
COMPANY_LOCATION = "Cebu City"
COMPANY_PHONE = "032-2670573"
COMPANY_MOBILE_SUN = "09325314857"
COMPANY_MOBILE_GLOBE = "09154657503"
SERVICES = "Sales    Installation    Service    Repair\nDuctworks"
MANAGER_DEFAULT = "J.B Yap Jr."
OUTPUT_DIR = Path(__file__).parent.parent / "data" / "output"

# Task templates with default costs
TASK_TEMPLATES = {
    "cleaning": {"name": "General cleaning", "cost": 400},
    "repair": {"name": "Repair", "cost": 3550},
    "installation": {"name": "Installation", "cost": 7500},
}

WARRANTY_OPTIONS = {
    "1": "None",
    "2": "Ninety (90) days, excluding compressor and all spare parts",
    "3": "Twelve (12) Months Only",
    "4": "Custom",
}

PAYMENT_OPTIONS = {
    "1": "COD (Cash on delivery)",
    "2": "50% Down payment, 50% Upon completion",
    "3": "70% Down Payment, 30% Cash Upon Completion",
    "4": "Cash Upon Completion",
    "5": "Cash Before Delivery + Installation",
    "6": "Full Payment Before Repair + Cleaning",
    "7": "Custom",
}

def print_header(text):
    """Print a section header."""
    print(f"\n{'━' * 60}")
    print(f"  {text}")
    print(f"{'━' * 60}\n")

def get_input(prompt, required=False, input_type=str, options=None):
    """Get validated user input."""
    while True:
        user_input = input(prompt).strip()

        if not user_input and required:
            print("  ⚠ This field is required. Please try again.")
            continue

        if not user_input:
            return None

        if options and user_input not in options:
            print(f"  ⚠ Invalid choice. Options: {', '.join(options.keys())}")
            continue

        if input_type == int:
            try:
                return int(user_input)
            except ValueError:
                print("  ⚠ Please enter a valid number.")
                continue

        if input_type == float:
            try:
                return float(user_input)
            except ValueError:
                print("  ⚠ Please enter a valid number.")
                continue

        return user_input

def get_yes_no(prompt):
    """Get yes/no response."""
    response = get_input(prompt, options={"y": "yes", "n": "no"})
    return response == "y"

def gather_header_info():
    """Gather customer and header information."""
    print_header("CUSTOMER INFORMATION")

    date_input = get_input("📅 Date (YYYY-MM-DD) [default: today]: ")
    try:
        date = datetime.strptime(date_input, "%Y-%m-%d") if date_input else datetime.today()
    except ValueError:
        print("  ⚠ Invalid date format. Using today.")
        date = datetime.today()

    customer_name = get_input("👤 Customer Name: ", required=True)
    customer_location = get_input("📍 Customer Location/Branch (optional): ")

    attention = None
    if get_yes_no("📬 Add Attention/Contact Person? (y/n): "):
        attention = get_input("   Contact Person Name: ", required=True)

    phone = None
    if get_yes_no("☎️  Add Phone Number? (y/n): "):
        phone = get_input("   Phone Number: ")

    installation_location = None
    if get_yes_no("📌 Add Installation Location Description? (y/n): "):
        installation_location = get_input("   Location Description: ")

    return {
        "date": date,
        "customer_name": customer_name,
        "customer_location": customer_location,
        "attention": attention,
        "phone": phone,
        "installation_location": installation_location,
    }

def select_document_type():
    """Select quotation document type."""
    print_header("DOCUMENT TYPE")
    print("  [1] Summary of Quotations (multiple items/departments)")
    print("  [2] Job to Be Done (single or multi-task job)\n")
    choice = get_input("Select type (1 or 2): ", required=True, options={"1": "summary", "2": "job"})
    return "summary" if choice == "1" else "job"

def gather_items(doc_type):
    """Gather quotation items with nested loops."""
    items = []
    item_num = 1

    while True:
        print_header(f"ITEM {item_num}")

        item_name = get_input(f"Item name (Department/Location/Job): ", required=True)
        ac_brand = get_input("AC Unit Brand (e.g., Panasonic, Daikin): ")
        ac_model = get_input("AC Unit Model/HP (e.g., 2.5Hp, 5Trs): ")

        tasks = []
        task_num = 1

        while True:
            print(f"\n  Task {task_num}:")
            print("    [1] General Cleaning  [2] Repair  [3] Installation  [4] Custom Task  [5] Done with tasks")
            task_choice = get_input("  Select (1-5): ", required=True, options={"1": "cleaning", "2": "repair", "3": "installation", "4": "custom", "5": "done"})

            if task_choice == "5":
                break

            if task_choice == "1":
                cost = get_input("    Cost (default ₱400): ", input_type=float)
                cost = cost if cost else 400.0
                tasks.append({"name": "General cleaning", "cost": cost, "quantity": 1})

            elif task_choice == "2":
                cost = get_input("    Cost (default ₱3,550): ", input_type=float)
                cost = cost if cost else 3550.0
                tasks.append({"name": "Repair", "cost": cost, "quantity": 1})

            elif task_choice == "3":
                cost = get_input("    Base Installation Cost (default ₱7,500): ", input_type=float)
                cost = cost if cost else 7500.0
                distance = get_input("    Distance in feet (0 if N/A): ", input_type=float)
                excess_distance = max(0, distance - 10) if distance else 0
                excess_cost = excess_distance * 350 if excess_distance > 0 else 0
                total_cost = cost + excess_cost
                tasks.append({
                    "name": f"Installation (base ₱{cost}, {distance}ft distance, excess ₱{excess_cost})",
                    "cost": total_cost,
                    "quantity": 1
                })

            elif task_choice == "4":
                task_name = get_input("    Task Description: ", required=True)
                cost = get_input("    Cost: ", required=True, input_type=float)
                quantity = get_input("    Quantity (default 1): ", input_type=float)
                quantity = quantity if quantity else 1
                tasks.append({"name": task_name, "cost": cost, "quantity": quantity})

            task_num += 1

        item_total = sum(t["cost"] * t["quantity"] for t in tasks)

        warranty = None
        if get_yes_no("  Add Warranty for this item? (y/n): "):
            print("    [1] None  [2] 90 days  [3] 12 Months  [4] Custom")
            w_choice = get_input("    Select (1-4): ", required=True, options={"1": "none", "2": "90", "3": "12", "4": "custom"})
            if w_choice == "4":
                warranty = get_input("    Warranty text: ", required=True)
            else:
                warranty = WARRANTY_OPTIONS[{"1": "1", "2": "2", "3": "3", "4": "4"}[w_choice]]

        items.append({
            "name": item_name,
            "ac_brand": ac_brand,
            "ac_model": ac_model,
            "tasks": tasks,
            "item_total": item_total,
            "warranty": warranty,
        })

        print(f"\n  ✓ Item {item_num} total: ₱{item_total:,.2f}")

        if not get_yes_no("\n  Add another item? (y/n): "):
            break

        item_num += 1

    return items

def gather_summary_info():
    """Gather overall quotation summary and footer info."""
    print_header("QUOTATION SUMMARY")

    print("  Warranty Options:")
    for k, v in WARRANTY_OPTIONS.items():
        print(f"    [{k}] {v}")
    warranty_choice = get_input("\n  Select warranty (1-4): ", required=True, options=WARRANTY_OPTIONS.keys())
    warranty = WARRANTY_OPTIONS[warranty_choice]
    if warranty_choice == "4":
        warranty = get_input("    Enter custom warranty: ", required=True)

    print("\n  Payment Terms Options:")
    for k, v in PAYMENT_OPTIONS.items():
        print(f"    [{k}] {v}")
    payment_choice = get_input("\n  Select payment terms (1-7): ", required=True, options=PAYMENT_OPTIONS.keys())
    payment = PAYMENT_OPTIONS[payment_choice]
    if payment_choice == "7":
        payment = get_input("    Enter custom payment terms: ", required=True)

    exceptions = None
    if get_yes_no("\n  Add exceptions? (y/n): "):
        exceptions = get_input("    Exceptions text: ", required=True)

    note = None
    if get_yes_no("\n  Add note (e.g., payee info)? (y/n): "):
        note = get_input("    Note text: ", required=True)

    manager = get_input(f"\n  Manager Name [default: {MANAGER_DEFAULT}]: ")
    manager = manager if manager else MANAGER_DEFAULT

    return {
        "warranty": warranty,
        "payment": payment,
        "exceptions": exceptions,
        "note": note,
        "manager": manager,
    }

def add_horizontal_line(doc):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def create_docx(header_info, doc_type, items, summary_info):
    """Create .docx file."""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Professional Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run(COMPANY_NAME)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Calibri'

    # Location
    location_para = doc.add_paragraph()
    location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location_para.paragraph_format.space_after = Pt(3)
    location_run = location_para.add_run(f"{COMPANY_LOCATION}")
    location_run.font.size = Pt(11)
    location_run.font.name = 'Calibri'

    # Horizontal line separator
    add_horizontal_line(doc)

    # Contact info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(2)
    contact = contact_para.add_run(f"Telephone: {COMPANY_PHONE} | Mobile: Sun-{COMPANY_MOBILE_SUN} | Globe-{COMPANY_MOBILE_GLOBE}")
    contact.font.size = Pt(9)
    contact.font.name = 'Calibri'

    # Services
    services_para = doc.add_paragraph()
    services_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    services_para.paragraph_format.space_after = Pt(12)
    services_run = services_para.add_run(SERVICES)
    services_run.font.size = Pt(10)
    services_run.font.name = 'Calibri'

    # Customer info section
    date_str = header_info["date"].strftime("%m/%d/%Y")

    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_after = Pt(3)
    date_label = date_para.add_run("Date: ")
    date_label.bold = True
    date_label.font.name = 'Calibri'
    date_val = date_para.add_run(date_str)
    date_val.font.name = 'Calibri'

    to_para = doc.add_paragraph()
    to_para.paragraph_format.space_after = Pt(2)
    to_label = to_para.add_run("To: ")
    to_label.bold = True
    to_label.font.name = 'Calibri'
    to_val = to_para.add_run(header_info["customer_name"])
    to_val.font.name = 'Calibri'

    if header_info["customer_location"]:
        loc_para = doc.add_paragraph(header_info["customer_location"])
        loc_para.paragraph_format.space_after = Pt(2)
        loc_para.paragraph_format.left_indent = Inches(0.25)
        for run in loc_para.runs:
            run.font.name = 'Calibri'

    if header_info["attention"]:
        att_para = doc.add_paragraph()
        att_para.paragraph_format.space_after = Pt(8)
        att_label = att_para.add_run("Attention: ")
        att_label.bold = True
        att_label.font.name = 'Calibri'
        att_text = header_info["attention"]
        if header_info["phone"]:
            att_text += f" | {header_info['phone']}"
        att_val = att_para.add_run(att_text)
        att_val.font.name = 'Calibri'

    doc.add_paragraph()  # Spacing

    # Salutation
    salutation = doc.add_paragraph("Sir/Madame,")
    salutation.paragraph_format.space_after = Pt(8)
    for run in salutation.runs:
        run.font.name = 'Calibri'

    # Opening
    opening = doc.add_paragraph()
    opening.paragraph_format.space_after = Pt(8)
    opening_text = ("This is with reference to your request for quotation for the installation/repair of the "
                    "following air-conditioner to be installed at:")
    opening_run = opening.add_run(opening_text)
    opening_run.font.name = 'Calibri'
    if header_info["installation_location"]:
        loc_run = opening.add_run(f" {header_info['installation_location']}")
        loc_run.font.name = 'Calibri'

    pleased = doc.add_paragraph("We are pleased to quote the following:")
    pleased.paragraph_format.space_after = Pt(8)
    for run in pleased.runs:
        run.font.name = 'Calibri'

    # Items section
    heading_text = "Summary of Quotations" if doc_type == "summary" else "Job to be done:"
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    heading_run = heading.add_run(heading_text)
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    heading_run.font.name = 'Calibri'

    grand_total = 0
    for i, item in enumerate(items, 1):
        # Item header
        item_header = doc.add_paragraph()
        item_header.paragraph_format.space_before = Pt(6)
        item_header.paragraph_format.space_after = Pt(3)
        item_header_run = item_header.add_run(f"{i}. {item['name']}")
        item_header_run.bold = True
        item_header_run.font.size = Pt(11)
        item_header_run.font.name = 'Calibri'
        if item["ac_brand"]:
            brand_run = item_header.add_run(f" - {item['ac_brand']}")
            brand_run.font.name = 'Calibri'
        if item["ac_model"]:
            model_run = item_header.add_run(f" {item['ac_model']}")
            model_run.font.name = 'Calibri'

        # Tasks
        for j, task in enumerate(item["tasks"], 1):
            task_para = doc.add_paragraph()
            task_para.paragraph_format.left_indent = Inches(0.25)
            task_para.paragraph_format.space_after = Pt(2)
            task_num = task_para.add_run(f"{j}. ")
            task_num.font.name = 'Calibri'
            task_name = task_para.add_run(f"{task['name']} – ₱{task['cost']:,.0f}")
            task_name.font.name = 'Calibri'
            if task["quantity"] > 1:
                qty_text = task_para.add_run(f" × {task['quantity']} = ₱{task['cost'] * task['quantity']:,.0f}")
                qty_text.font.name = 'Calibri'

        # Item total
        total_para = doc.add_paragraph()
        total_para.paragraph_format.left_indent = Inches(0.25)
        total_para.paragraph_format.space_after = Pt(3)
        total_run = total_para.add_run(f"Total Price – ₱ {item['item_total']:,.2f}")
        total_run.bold = True
        total_run.font.name = 'Calibri'

        # Item warranty
        if item.get("warranty"):
            warranty_para = doc.add_paragraph()
            warranty_para.paragraph_format.left_indent = Inches(0.25)
            warranty_para.paragraph_format.space_after = Pt(8)
            warranty_run = warranty_para.add_run(f"Warranty: {item['warranty']}")
            warranty_run.font.name = 'Calibri'

        grand_total += item['item_total']

    # Grand total
    if len(items) > 1:
        grand_para = doc.add_paragraph()
        grand_para.paragraph_format.space_before = Pt(8)
        grand_para.paragraph_format.space_after = Pt(12)
        grand_run = grand_para.add_run(f"Total Price of All Items – ₱ {grand_total:,.2f}")
        grand_run.bold = True
        grand_run.font.name = 'Calibri'

    # Summary info section
    add_horizontal_line(doc)

    if summary_info.get("note"):
        note_para = doc.add_paragraph()
        note_para.paragraph_format.space_after = Pt(6)
        note_label = note_para.add_run("Note: ")
        note_label.bold = True
        note_label.font.name = 'Calibri'
        note_val = note_para.add_run(summary_info['note'])
        note_val.font.name = 'Calibri'

    payment_para = doc.add_paragraph()
    payment_para.paragraph_format.space_after = Pt(3)
    payment_label = payment_para.add_run("Terms of Payment: ")
    payment_label.bold = True
    payment_label.font.name = 'Calibri'
    payment_val = payment_para.add_run(summary_info['payment'])
    payment_val.font.name = 'Calibri'

    warranty_para = doc.add_paragraph()
    warranty_para.paragraph_format.space_after = Pt(3)
    warranty_label = warranty_para.add_run("Warranty: ")
    warranty_label.bold = True
    warranty_label.font.name = 'Calibri'
    warranty_val = warranty_para.add_run(summary_info['warranty'])
    warranty_val.font.name = 'Calibri'

    if summary_info.get("exceptions"):
        excep_para = doc.add_paragraph()
        excep_para.paragraph_format.space_after = Pt(12)
        excep_label = excep_para.add_run("Exception: ")
        excep_label.bold = True
        excep_label.font.name = 'Calibri'
        excep_val = excep_para.add_run(summary_info['exceptions'])
        excep_val.font.name = 'Calibri'

    # Closing
    thanks_para = doc.add_paragraph()
    thanks_para.paragraph_format.space_before = Pt(6)
    thanks_para.paragraph_format.space_after = Pt(12)
    thanks_run = thanks_para.add_run("Thank you very much for giving us the opportunity to quote and we hope to have the pleasure of serving you.")
    thanks_run.font.name = 'Calibri'

    doc.add_paragraph()  # Spacing

    closing = doc.add_paragraph("Very Truly Yours,")
    closing.paragraph_format.space_after = Pt(40)
    for run in closing.runs:
        run.font.name = 'Calibri'

    # Signature block
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = False

    # Conforme line
    conforme_cell = sig_table.rows[0].cells[0]
    conforme_para = conforme_cell.paragraphs[0]
    conforme_run = conforme_para.add_run("Conforme:_______________")
    conforme_run.font.name = 'Calibri'

    # Manager name cell
    manager_cell = sig_table.rows[0].cells[1]
    manager_para = manager_cell.paragraphs[0]
    manager_name_run = manager_para.add_run(summary_info['manager'])
    manager_name_run.font.name = 'Calibri'
    manager_name_run.bold = True
    manager_para.add_run("\n")
    manager_title_run = manager_para.add_run("Manager")
    manager_title_run.font.name = 'Calibri'

    # Date line
    date_cell = sig_table.rows[1].cells[0]
    date_para = date_cell.paragraphs[0]
    date_run = date_para.add_run("Date:_______________")
    date_run.font.name = 'Calibri'

    # Empty cell for second column
    sig_table.rows[1].cells[1].text = ""

    return doc

def save_and_convert(doc, filename):
    """Save .docx and convert to .pdf."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    docx_path = OUTPUT_DIR / f"{filename}.docx"
    pdf_path = OUTPUT_DIR / f"{filename}.pdf"

    # Save docx
    doc.save(docx_path)
    print(f"  ✓ Saved: {docx_path}")

    # Convert to PDF
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTPUT_DIR), str(docx_path)],
            check=True,
            capture_output=True
        )
        print(f"  ✓ Converted: {pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"  ⚠ PDF conversion failed: {e.stderr.decode()}")
        print("  (DocX file was saved successfully)")
    except FileNotFoundError:
        print("  ⚠ LibreOffice not found. PDF conversion skipped.")
        print("  (DocX file was saved successfully)")

def main():
    """Main execution flow."""
    print("\n" + "=" * 60)
    print("  QUOTATION GENERATOR - Cebu Best Value Trading Corp.")
    print("=" * 60)

    try:
        # Gather info
        header_info = gather_header_info()
        doc_type = select_document_type()
        items = gather_items(doc_type)
        summary_info = gather_summary_info()

        # Preview
        print_header("PREVIEW")
        print(f"Customer: {header_info['customer_name']}")
        print(f"Date: {header_info['date'].strftime('%m/%d/%Y')}")
        print(f"Type: {doc_type.upper()}")
        print(f"Items: {len(items)}")
        grand_total = sum(item['item_total'] for item in items)
        print(f"Total: ₱{grand_total:,.2f}")

        if not get_yes_no("\n✓ Save & Generate files? (y/n): "):
            print("\n✗ Cancelled.")
            return

        # Generate files
        print_header("GENERATING FILES")
        doc = create_docx(header_info, doc_type, items, summary_info)
        filename = f"{header_info['customer_name']}-{header_info['date'].strftime('%m_%d_%Y')}"
        save_and_convert(doc, filename)

        print(f"\n✓ Files saved to: {OUTPUT_DIR}\n")

    except KeyboardInterrupt:
        print("\n\n✗ Cancelled by user.")
        sys.exit(0)
    except Exception as e:
        print(f"\n✗ Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
